"""Extract and chunk text from uploaded process documents (PDF/DOCX/TXT/MD)
so they can be ingested into the same knowledge graph as videos -- process
intake forms, architecture checklists, SOPs, and the like.

Deliberately reuses vidblog's exact sections.json schema (video_title,
channel, url, sections[{index,start,end,raw_transcript}]) with a "kind":
"document" marker, so kgwiki's graph_builder / search / RAG code needs no
separate code path for documents -- they're just sources with no screenshots
and no real timeline.
"""
from __future__ import annotations

import hashlib
import json
import os
import re
import shutil

_ALLOWED_EXTS = {".pdf", ".docx", ".txt", ".md"}


def _slugify_id(path: str) -> str:
    base = os.path.splitext(os.path.basename(path))[0]
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", base).strip("-").lower()[:40] or "document"
    digest = hashlib.sha1(os.path.abspath(path).encode("utf-8")).hexdigest()[:8]
    return f"{slug}-{digest}"


def is_supported_document(filename: str) -> bool:
    return os.path.splitext(filename)[1].lower() in _ALLOWED_EXTS


def _extract_pdf(path: str) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
            # Intake forms/checklists are often tabular -- extract_text() alone
            # tends to mangle columns, so pull tables out explicitly too.
            for table in page.extract_tables() or []:
                rows = [" | ".join(cell or "" for cell in row) for row in table]
                if rows:
                    parts.append("\n".join(rows))
    return "\n\n".join(parts)


def _extract_docx(path: str) -> str:
    import docx

    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n\n".join(parts)


def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext in (".txt", ".md"):
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    raise ValueError(f"Unsupported document type: {ext}")


def chunk_text(text: str, words_per_chunk: int = 180) -> list[str]:
    """Paragraph-aware chunking: keeps paragraphs intact where possible,
    grouping them up to a target size rather than cutting mid-sentence."""
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    for para in paragraphs:
        words = para.split()
        current.append(para)
        current_len += len(words)
        if current_len >= words_per_chunk:
            chunks.append("\n\n".join(current))
            current = []
            current_len = 0
    if current:
        chunks.append("\n\n".join(current))
    if not chunks and text.strip():
        chunks = [text.strip()]
    return chunks


def _heading_from_chunk(chunk: str, index: int) -> str:
    first_line = chunk.strip().splitlines()[0].strip()
    if 3 <= len(first_line) <= 80:
        return first_line.rstrip(":.")
    words = chunk.split()[:7]
    return " ".join(words).strip(" ,.") or f"Section {index + 1}"


def set_folder(work_dir: str, folder_id: str | None) -> None:
    meta_path = os.path.join(work_dir, "meta.json")
    meta = {}
    if os.path.exists(meta_path):
        try:
            with open(meta_path, "r", encoding="utf-8") as f:
                meta = json.load(f)
        except (json.JSONDecodeError, OSError):
            meta = {}
    meta["folder_id"] = folder_id
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(meta, f)


def get_folder(work_dir: str) -> str | None:
    meta_path = os.path.join(work_dir, "meta.json")
    if not os.path.exists(meta_path):
        return None
    try:
        with open(meta_path, "r", encoding="utf-8") as f:
            return json.load(f).get("folder_id")
    except (json.JSONDecodeError, OSError):
        return None


def ingest_document(
    path: str,
    out_root: str,
    title: str | None = None,
    folder_id: str | None = None,
) -> str:
    """Extract, chunk, and write sections.json for a document. Returns the
    doc_id (its work_dir name under out_root)."""
    if not is_supported_document(path):
        raise ValueError(f"Unsupported document type: {os.path.splitext(path)[1]}")

    doc_id = _slugify_id(path)
    work_dir = os.path.join(out_root, doc_id)
    os.makedirs(work_dir, exist_ok=True)

    ext = os.path.splitext(path)[1].lower()
    dest_path = os.path.join(work_dir, f"document{ext}")
    if not os.path.exists(dest_path):
        try:
            os.link(path, dest_path)  # instant, no extra disk, if same volume
        except OSError:
            shutil.copy2(path, dest_path)

    text = extract_text(dest_path)
    if not text.strip():
        raise ValueError("No extractable text found in this document (it may be a scanned image PDF).")

    chunks = chunk_text(text)
    display_title = title or os.path.splitext(os.path.basename(path))[0].strip()

    sections = []
    for i, chunk in enumerate(chunks):
        # Sequential placeholder "timestamps", purely to give sections a
        # stable sort order -- documents have no real timeline, and the UI
        # hides time-based fields for kind="document" sources.
        m, s = divmod(i, 60)
        sections.append(
            {
                "index": i,
                "start": f"{m}:{s:02d}",
                "end": f"{m}:{s:02d}",
                "raw_transcript": chunk,
                "heading": _heading_from_chunk(chunk, i),
            }
        )

    payload = {
        "video_title": display_title,
        "channel": "Document",
        "url": "",
        "kind": "document",
        "source_filename": os.path.basename(path),
        "sections": sections,
    }
    with open(os.path.join(work_dir, "sections.json"), "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)

    if folder_id:
        set_folder(work_dir, folder_id)

    return doc_id
